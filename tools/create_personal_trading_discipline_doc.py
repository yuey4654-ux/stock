from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(r"C:\Users\Administrator\Documents\股票分析\个人交易纪律表_V1.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx] / 567)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_p(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        set_east_asia_font(r)
        rest = text[len(bold_prefix):]
        if rest:
            r2 = p.add_run(rest)
            set_east_asia_font(r2)
    else:
        r = p.add_run(text)
        set_east_asia_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_east_asia_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_east_asia_font(r)


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_east_asia_font(r)
    return p


def fill_row(row, values, header=False):
    for cell, value in zip(row.cells, values):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(value)
        set_east_asia_font(run)
        if header:
            run.bold = True
            set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    fill_row(table.rows[0], headers, header=True)
    for row_values in rows:
        row = table.add_row()
        fill_row(row, row_values)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def set_document_styles(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_doc():
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("个人交易纪律表 V1")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("适用于 50 万左右资金、A 股 + 港股短线波段、持仓几天至两周的交易纪律")
    set_east_asia_font(r)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")

    add_heading(doc, "一、我的交易定位", 1)
    add_p(doc, "本纪律表不是预测工具，而是减少随机交易、控制亏损、提高复盘质量的执行系统。第一阶段目标不是重仓赚快钱，而是把“听推荐就买、热门票追高、港股越跌越补”的习惯压下来。")
    add_table(
        doc,
        ["项目", "当前设定"],
        [
            ["总资金", "约 50 万元"],
            ["当前仓位", "约 6 成；短线新手阶段先不主动提高仓位"],
            ["主要市场", "A 股约 4 成，港股约 2 成，美股少量可暂不纳入纪律核心"],
            ["持仓周期", "几天到两周，以短线波段为主"],
            ["核心问题", "容易被推荐票和热门票影响，买在利好结束阶段"],
            ["阶段目标", "先做到每笔交易有理由、有止损、有计划、有复盘"],
        ],
        [1900, 7460],
    )

    add_heading(doc, "二、总纪律：先活下来，再提高胜率", 1)
    add_bullets(
        doc,
        [
            "别人推荐的票，不能当天无脑买；必须先进观察池。",
            "连续上涨 3-5 天后才看到的利好，默认已经进入后半段，禁止追高。",
            "买入前必须写清楚：为什么买、哪里错、错了亏多少、涨了怎么卖。",
            "单票首次买入不超过 2 万元；很有把握也不超过 5 万元。",
            "亏损连续 3 笔后，暂停新开仓 3 天，只复盘不交易。",
            "港股做 T 只能用于增强收益，不能用于拯救亏损仓位。",
        ],
    )

    add_heading(doc, "三、仓位与亏损上限", 1)
    add_table(
        doc,
        ["场景", "最大买入金额", "常规止损", "灾难线", "动作"],
        [
            ["普通试错", "≤ 20,000", "亏 6%-8%", "亏 10%", "跌破计划止损直接走"],
            ["较高把握", "20,000-50,000", "亏 5%-6%", "亏 8%", "先半仓，确认后再加"],
            ["港股 T 仓", "不超过该票计划仓位 1/3", "日内不对立即撤", "亏 5%", "T 仓不能变成长线仓"],
            ["总仓位", "新手阶段 4-6 成", "账户回撤 5% 降仓", "账户回撤 8%-10%", "停止新开仓，复盘持仓"],
        ],
        [1600, 1900, 1900, 1600, 2360],
    )

    add_heading(doc, "四、只做两类买点", 1)
    add_heading(doc, "模型 A：强势回踩", 2)
    add_bullets(
        doc,
        [
            "股票近期明显强于大盘，且处在热门板块或强趋势中。",
            "股价在 20 日线或 60 日线上方，前期有放量上涨。",
            "回调时缩量，回踩 5 日线、10 日线、20 日线或前高支撑不破。",
            "出现重新放量转强的阳线后再买，不因为“跌了”就提前买。",
            "止损位放在支撑位下方，或按 6%-8% 的常规止损执行。",
        ],
    )
    add_heading(doc, "模型 B：平台突破", 2)
    add_bullets(
        doc,
        [
            "横盘至少 2-4 周，平台高点清晰。",
            "突破当日成交量明显放大，大盘环境不能太弱。",
            "突破后不能离 20 日线太远，避免追在短线情绪高潮。",
            "跌回突破位且收不回，必须减仓或止损。",
        ],
    )

    add_heading(doc, "五、禁止交易清单", 1)
    add_table(
        doc,
        ["禁止事项", "原因", "替代动作"],
        [
            ["听到推荐马上买", "容易接利好最后一棒", "先放入观察池，等回踩或突破确认"],
            ["高开 5% 以上冲进去", "赔率差，止损距离大", "等盘中回落承接，或次日再看"],
            ["亏损中不断做 T 摊低", "会把短线错单变成长线套牢", "先确认是否跌破逻辑，必要时止损"],
            ["没有止损价就买", "盘中会临时改剧本", "先写交易卡，再下单"],
            ["仓位超过计划后继续补", "单笔错误会伤到账户", "停止加仓，等复盘结论"],
            ["连续亏损还想翻本", "情绪交易概率大", "暂停 3 天，只观察不交易"],
        ],
        [2200, 3600, 3560],
    )

    add_heading(doc, "六、每日复盘 20 分钟", 1)
    add_p(doc, "每天只复盘 5 件事。不要追求复杂，先做到连续记录 20 个交易日。")
    add_table(
        doc,
        ["复盘问题", "填写区"],
        [
            ["1. 今天大盘环境：强 / 中性 / 弱", ""],
            ["2. 今天最强板块是哪 2-3 个？", ""],
            ["3. 我的持仓有没有跌破买入逻辑？", ""],
            ["4. 今天有没有冲动交易？如果有，原因是什么？", ""],
            ["5. 明天只观察哪 3 只票？", ""],
        ],
        [3600, 5760],
    )

    add_heading(doc, "七、交易前检查表", 1)
    add_p(doc, "下单前必须至少满足 7 项中的 5 项；如果第 1、2、6 项任意一项不满足，禁止买入。")
    add_table(
        doc,
        ["检查项", "是否通过"],
        [
            ["□ 我知道这只票为什么涨，且不是刚看到利好才追进去", ""],
            ["□ 它属于强势回踩或平台突破，不是随手买", ""],
            ["□ 大盘环境不差，至少不是明显杀跌日", ""],
            ["□ 买入价距离支撑位不远，止损空间可控", ""],
            ["□ 首次买入金额不超过 2 万，或符合较高把握规则", ""],
            ["□ 已写好止损价和止盈计划", ""],
            ["□ 如果明天低开，我知道该怎么处理", ""],
        ],
        [7200, 2160],
    )

    add_heading(doc, "八、单笔交易卡", 1)
    add_table(
        doc,
        ["字段", "填写"],
        [
            ["股票代码 / 名称", ""],
            ["买入日期 / 买入价格", ""],
            ["买入金额 / 计划仓位", ""],
            ["交易模式", "强势回踩 / 平台突破 / 其他"],
            ["买入理由", "1. \n2. \n3. "],
            ["止损价格 / 止损条件", ""],
            ["第一目标价 / 止盈动作", ""],
            ["计划持有时间", "几天 / 一周 / 两周"],
            ["如果明天跌了怎么办", ""],
            ["如果明天涨了怎么办", ""],
            ["复盘结论", "按系统交易 / 情绪交易 / 需要改进"],
        ],
        [2500, 6860],
    )

    add_heading(doc, "九、当前持仓分类表", 1)
    add_p(doc, "把现有持仓先分成四类。C 类优先处理，D 类禁止加仓。")
    add_table(
        doc,
        ["股票", "仓位金额", "当前盈亏", "分类", "处理动作"],
        [
            ["", "", "", "A 盈利且趋势仍在", "保留，跌破 20 日线或关键支撑减仓"],
            ["", "", "", "B 小亏但逻辑仍在", "设明确止损，不加仓"],
            ["", "", "", "C 亏损扩大且理由说不清", "优先减仓或止损"],
            ["", "", "", "D 纯听推荐买入", "禁止加仓，等待反弹处理"],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
        ],
        [1500, 1600, 1500, 2500, 2260],
    )

    add_heading(doc, "十、每周复盘统计", 1)
    add_table(
        doc,
        ["指标", "本周记录"],
        [
            ["本周交易次数", ""],
            ["盈利笔数 / 亏损笔数", ""],
            ["最大单笔亏损", ""],
            ["是否有无计划交易", ""],
            ["是否有追高推荐票", ""],
            ["是否遵守止损", ""],
            ["下周只改进一件事", ""],
        ],
        [3000, 6360],
    )

    add_heading(doc, "十一、执行口令", 1)
    add_numbered(
        doc,
        [
            "推荐票先进观察池，不直接买。",
            "只做强势回踩和平台突破。",
            "小仓试错，错了快走。",
            "港股做 T 不救亏损仓。",
            "每天复盘 20 分钟，连续记录 20 个交易日。",
        ],
    )

    doc.core_properties.title = "个人交易纪律表 V1"
    doc.core_properties.subject = "短线交易纪律、仓位控制、复盘模板"
    doc.core_properties.author = "Codex"
    doc.save(OUT)


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc()
    print(OUT)
