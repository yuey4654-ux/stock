from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path(r"C:\Users\Administrator\Documents\股票分析\个人交易纪律表_V1.docx")
OUT = Path(r"C:\Users\Administrator\Documents\股票分析\个人交易纪律表_V2_加入选股指标.docx")


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_east_asia_font(r)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_east_asia_font(r)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx] / 567)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def fill_row(row, values, header=False):
    for cell, value in zip(row.cells, values):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(value)
        set_east_asia_font(run)
        if header:
            run.bold = True
            set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    fill_row(table.rows[0], headers, header=True)
    for row_values in rows:
        fill_row(table.add_row(), row_values)
    set_table_width(table, widths)
    doc.add_paragraph()


def main():
    copyfile(SRC, OUT)
    doc = Document(OUT)

    doc.add_page_break()
    add_heading(doc, "十二、选股指标组合 V1：先过滤，再买入", 1)
    add_p(doc, "本节根据 AlphaGBM 的趋势、情绪、资金、估值框架，并结合公开炒股平台常见短线共识，提炼为可执行指标组合。回测使用过去一年公开行情的小样本验证，结论用于约束交易，不代表未来必然盈利。")

    add_heading(doc, "1. 平台和研究共识", 2)
    add_bullets(
        doc,
        [
            "单一指标胜率不稳定，必须把趋势、量能、位置、情绪和风控组合起来。",
            "短线更适合做强势股回踩后的再转强，而不是利好公开后追高。",
            "MACD、RSI 适合作为过滤器，不适合作为唯一买点。",
            "成交量、换手率、均线结构和近几日涨幅，是避免买在利好尾声的关键过滤项。",
        ],
    )

    add_heading(doc, "2. 回测摘要", 2)
    add_table(
        doc,
        ["指标组合", "交易次数", "胜率", "平均收益", "盈亏比/利润因子", "结论"],
        [
            ["平台突破", "9", "44.4%", "-0.98%", "0.61", "单独使用不适合新手追高"],
            ["强势回踩", "12", "58.3%", "+2.17%", "2.04", "可作为主模型"],
            ["MACD+RSI趋势过滤", "1", "0.0%", "-6.20%", "0.00", "信号太少，不单独使用"],
            ["综合高约束组合", "8", "62.5%", "+2.91%", "2.84", "当前优先采用"],
            ["综合组合参数优化", "8", "75.0%", "+7.01%", "9.19", "样本偏小，只作为观察上限"],
        ],
        [1900, 1200, 1000, 1200, 1600, 2460],
    )

    add_heading(doc, "3. 我的优先选股组合", 2)
    add_p(doc, "只在以下条件多数满足时，才允许把股票加入候选池；真正买入仍要等待交易卡确认。")
    add_table(
        doc,
        ["维度", "硬条件", "目的"],
        [
            ["趋势", "MA10 > MA20 > MA60，股价在 MA20 和 MA60 上方", "只做向上的票"],
            ["位置", "过去 20 日涨幅 8%-35%，且不是连续暴涨后第一次看到消息", "避免利好末端追高"],
            ["买点", "回踩 MA10/MA20 或前高支撑不破，次日重新放量转强", "等低风险买点"],
            ["量能", "买入日成交量 ≥ 20 日均量，最好 1.2-2.0 倍", "确认资金仍在"],
            ["RSI", "RSI14 在 45-70 之间", "避开过弱和过热"],
            ["MACD", "DIF > DEA，或即将金叉但不能明显背离", "只做趋势确认"],
            ["风险", "买入价到止损位不超过 6%-8%", "控制单笔亏损"],
        ],
        [1300, 4300, 3760],
    )

    add_heading(doc, "4. 每日选股打分表", 2)
    add_table(
        doc,
        ["检查项", "分值", "是否满足"],
        [
            ["MA10 > MA20 > MA60", "20", ""],
            ["股价站上 MA20 和 MA60", "15", ""],
            ["回踩支撑不破后重新转强", "20", ""],
            ["成交量 ≥ 20 日均量", "15", ""],
            ["RSI14 在 45-70", "10", ""],
            ["近 5 日涨幅 < 18%，未明显过热", "10", ""],
            ["买入到止损距离 ≤ 8%", "10", ""],
        ],
        [6200, 1100, 2060],
    )
    add_p(doc, "执行规则：总分 80 分以上才可进入交易计划；70-79 分只观察；低于 70 分禁止买入。")

    add_heading(doc, "5. 选股红线", 2)
    add_bullets(
        doc,
        [
            "连续大涨后才被推荐，禁止当天追。",
            "RSI 超过 75 且放巨量长上影，禁止追。",
            "股价距离 MA20 超过 15%，禁止追。",
            "成交量突然放大但收盘走弱，禁止买入。",
            "止损空间超过 8%，即使逻辑再好也不买。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
